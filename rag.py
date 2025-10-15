import os
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.schema import Document
from openai import OpenAI
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document

load_dotenv()

class KnowledgeBase:
    def __init__(self):
        try:
            self.embeddings = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
            print("✓ Embeddings model loaded successfully")
        except Exception as e:
            print(f"Warning: Could not load embeddings model: {e}")
            # Fallback to a simpler approach
            self.embeddings = None
        self.vectorstore = None
        self.client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

    def add_documents(self, file_paths):
        documents = []
        for file_path in file_paths:
            if file_path.endswith('.pdf'):
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            else:
                continue

        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        docs = text_splitter.split_documents(documents)

        if self.embeddings is None:
            raise Exception("Embeddings model not available. Please check your internet connection and try again.")

        if self.vectorstore is None:
            self.vectorstore = FAISS.from_documents(docs, self.embeddings)
        else:
            self.vectorstore.add_documents(docs)

    def query(self, question):
        if self.vectorstore is None:
            return "No documents uploaded yet."
        retriever = self.vectorstore.as_retriever()
        docs = retriever.invoke(question)
        context = " ".join([d.page_content for d in docs])
        prompt = f"Using these documents, answer the user’s question succinctly.\n\nDocuments: {context}\n\nQuestion: {question}"
        response = self.client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=150
        )
        return response.choices[0].message.content
