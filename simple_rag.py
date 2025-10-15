import os
from langchain.text_splitter import CharacterTextSplitter
from langchain.schema import Document
from openai import OpenAI
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document as DocxDocument

load_dotenv()

class SimpleKnowledgeBase:
    def __init__(self):
        self.documents = []  # Store documents as simple text
        self.client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        print("✓ Simple Knowledge Base initialized")

    def add_documents(self, file_paths):
        documents = []
        for file_path in file_paths:
            print(f"Processing file: {file_path}")
            if file_path.endswith('.pdf'):
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            elif file_path.endswith('.docx'):
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                documents.append(Document(page_content=text, metadata={"source": file_path}))
            else:
                print(f"Skipping unsupported file: {file_path}")
                continue

        # Store documents for simple retrieval
        self.documents.extend(documents)
        print(f"✓ Added {len(documents)} documents to knowledge base")

    def query(self, question):
        if not self.documents:
            return "No documents uploaded yet."
        
        # Simple retrieval - get all documents (can be improved later)
        context = " ".join([doc.page_content for doc in self.documents])
        
        # Truncate context if too long
        if len(context) > 4000:
            context = context[:4000] + "..."
        
        prompt = f"""Using the following documents, answer the user's question concisely and accurately.

Documents: {context}

Question: {question}

Answer:"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=300,
                temperature=0.3
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating answer: {str(e)}"
